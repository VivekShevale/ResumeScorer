"""
utils/file_parser.py
--------------------
Extracts plain text from:
  - PDF  (.pdf)  — via pdfplumber (primary) + pymupdf (fallback)
                   Also extracts hyperlinks from PDF annotations
  - DOCX (.docx) — via python-docx
  - DOC  (.doc)  — reads as binary text (best-effort)
"""

from __future__ import annotations
import re
from pathlib import Path
from utils.logger import get_logger

logger = get_logger(__name__)


def _extract_pdf_hyperlinks(file_path: str) -> str:
    """
    Extract hyperlinks embedded in PDF annotations.
    Many modern resumes (Canva, Enhancv, Novoresume) store GitHub/LinkedIn/LeetCode
    as clickable links in annotations — not as visible text.
    Returns a newline-joined string of all found social/profile URLs.
    """
    links = []
    try:
        import pdfplumber
        with pdfplumber.open(file_path) as pdf:
            for page in pdf.pages:
                annots = page.annots or []
                for annot in annots:
                    uri = annot.get("uri")
                    if uri and isinstance(uri, (str, bytes)):
                        if isinstance(uri, bytes):
                            uri = uri.decode("utf-8", errors="ignore")
                        uri = uri.strip()
                        # Only keep social/professional profile links
                        if any(domain in uri.lower() for domain in [
                            "github.com", "linkedin.com", "leetcode.com",
                            "codeforces.com", "codechef.com", "twitter.com",
                            "x.com", "kaggle.com", "stackoverflow.com",
                        ]):
                            if uri not in links:
                                links.append(uri)
    except Exception as e:
        logger.debug(f"PDF hyperlink extraction failed: {e}")

    if links:
        logger.info(f"PDF hyperlinks extracted: {links}")
    return "\n".join(links)


def extract_text_from_pdf(file_path: str) -> tuple[str, str | None]:
    """
    Extract text from PDF using pdfplumber; fallback to pymupdf.
    Appends any hyperlinks found in PDF annotations as [EXTRACTED_LINKS] section.
    """
    text = ""
    try:
        import pdfplumber
        with pdfplumber.open(file_path) as pdf:
            pages = []
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    pages.append(page_text)
            text = "\n".join(pages)
        if text.strip():
            logger.info(f"PDF parsed via pdfplumber | chars={len(text)} | file={Path(file_path).name}")
            hyperlinks = _extract_pdf_hyperlinks(file_path)
            if hyperlinks:
                text = text + "\n\n[EXTRACTED_LINKS]\n" + hyperlinks
            return text, None
    except Exception as e:
        logger.warning(f"pdfplumber failed: {e}, trying pymupdf fallback")

    # -- Fallback: pymupdf --
    try:
        import fitz  # pymupdf
        doc = fitz.open(file_path)
        pages = [doc[i].get_text() for i in range(len(doc))]
        text = "\n".join(pages)
        logger.info(f"PDF parsed via pymupdf | chars={len(text)} | file={Path(file_path).name}")
        hyperlinks = _extract_pdf_hyperlinks(file_path)
        if hyperlinks:
            text = text + "\n\n[EXTRACTED_LINKS]\n" + hyperlinks
        return text, None
    except Exception as e:
        msg = f"Both PDF parsers failed: {e}"
        logger.error(msg)
        return "", msg


def extract_text_from_docx(file_path: str) -> tuple[str, str | None]:
    """Extract text from DOCX using python-docx."""
    try:
        from docx import Document
        doc = Document(file_path)
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        # Also extract table cells
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        paragraphs.append(cell.text.strip())
        text = "\n".join(paragraphs)
        logger.info(f"DOCX parsed | paragraphs={len(paragraphs)} | chars={len(text)} | file={Path(file_path).name}")
        return text, None
    except Exception as e:
        msg = f"DOCX parsing failed: {e}"
        logger.error(msg)
        return "", msg


def extract_text_from_doc(file_path: str) -> tuple[str, str | None]:
    """
    Best-effort .doc extraction.
    Falls back to reading raw bytes and decoding printable chars.
    """
    try:
        # Try treating it as docx first (some .doc files are actually docx)
        text, err = extract_text_from_docx(file_path)
        if text.strip():
            return text, None
    except Exception:
        pass

    try:
        with open(file_path, "rb") as f:
            raw = f.read()
        text = raw.decode("latin-1", errors="ignore")
        printable = "".join(c if c.isprintable() or c in "\n\t" else " " for c in text)
        printable = re.sub(r" {3,}", " ", printable)
        printable = re.sub(r"\n{3,}", "\n\n", printable)
        logger.warning(f".doc raw fallback used | chars={len(printable)} | file={Path(file_path).name}")
        return printable, None
    except Exception as e:
        msg = f".doc parsing failed: {e}"
        logger.error(msg)
        return "", msg


def extract_text(file_path: str) -> tuple[str, str | None]:
    """
    Main entry point. Detects file type and calls appropriate parser.
    Returns (text, error_message_or_None)
    """
    path = Path(file_path)
    if not path.exists():
        return "", f"File not found: {file_path}"

    suffix = path.suffix.lower()

    if suffix == ".pdf":
        return extract_text_from_pdf(file_path)
    elif suffix == ".docx":
        return extract_text_from_docx(file_path)
    elif suffix == ".doc":
        return extract_text_from_doc(file_path)
    else:
        return "", f"Unsupported file type: {suffix}. Supported: .pdf, .docx, .doc"